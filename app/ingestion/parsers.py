"""Document parsers.

Pluggable backends behind one interface. Docling gives layout-aware output
(tables, reading order, headings) which matters enormously for banking policy
PDFs; the native parsers are the dependency-light fallback.
"""
from __future__ import annotations

import io
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from typing import Any

from app.core.config import settings
from app.core.logging import get_logger

log = get_logger(__name__)


@dataclass(slots=True)
class ParsedBlock:
    text: str
    page_number: int | None = None
    heading: str | None = None
    section_path: str | None = None
    block_type: str = "paragraph"  # paragraph | table | list | heading | caption
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(slots=True)
class ParsedDocument:
    blocks: list[ParsedBlock]
    page_count: int
    backend: str
    title: str | None = None
    language: str = "en"
    metadata: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(b.text for b in self.blocks if b.text.strip())


class BaseParser(ABC):
    name: str = "base"
    supported: tuple[str, ...] = ()

    @abstractmethod
    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument: ...

    def handles(self, content_type: str, filename: str) -> bool:
        return content_type in self.supported or filename.lower().endswith(
            tuple(f".{s.rsplit('/', 1)[-1]}" for s in self.supported)
        )


class PdfParser(BaseParser):
    name = "pypdf"
    supported = ("application/pdf",)

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        blocks: list[ParsedBlock] = []
        for page_no, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(ParsedBlock(text=text, page_number=page_no))
        title = None
        if reader.metadata:
            title = reader.metadata.get("/Title") or None
        return ParsedDocument(
            blocks=blocks, page_count=len(reader.pages), backend=self.name, title=title
        )


class DocxParser(BaseParser):
    name = "python-docx"
    supported = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        import docx  # type: ignore

        document = docx.Document(io.BytesIO(data))
        blocks: list[ParsedBlock] = []
        heading_stack: list[str] = []

        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower()
            if style.startswith("heading"):
                try:
                    level = int(style.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                heading_stack = heading_stack[: level - 1] + [text]
                blocks.append(
                    ParsedBlock(
                        text=text, heading=text, block_type="heading",
                        section_path=" > ".join(heading_stack),
                    )
                )
            else:
                blocks.append(
                    ParsedBlock(
                        text=text,
                        heading=heading_stack[-1] if heading_stack else None,
                        section_path=" > ".join(heading_stack) or None,
                    )
                )

        for table in document.tables:
            rows = [
                " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
            ]
            if rows:
                blocks.append(
                    ParsedBlock(
                        text="\n".join(rows), block_type="table",
                        section_path=" > ".join(heading_stack) or None,
                    )
                )

        return ParsedDocument(blocks=blocks, page_count=1, backend=self.name)


class TextParser(BaseParser):
    name = "text"
    supported = ("text/plain", "text/markdown", "text/csv", "application/json")

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        text = data.decode("utf-8", errors="replace")
        blocks = [
            ParsedBlock(text=part.strip())
            for part in text.split("\n\n")
            if part.strip()
        ]
        return ParsedDocument(blocks=blocks, page_count=1, backend=self.name)


class HtmlParser(BaseParser):
    name = "bs4"
    supported = ("text/html", "application/xhtml+xml")

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(data.decode("utf-8", errors="replace"), "html.parser")
        for tag in soup(["script", "style", "nav", "footer"]):
            tag.decompose()
        blocks: list[ParsedBlock] = []
        heading: str | None = None
        for element in soup.find_all(["h1", "h2", "h3", "p", "li", "table"]):
            text = element.get_text(" ", strip=True)
            if not text:
                continue
            if element.name in {"h1", "h2", "h3"}:
                heading = text
                blocks.append(ParsedBlock(text=text, heading=text, block_type="heading"))
            else:
                blocks.append(
                    ParsedBlock(
                        text=text, heading=heading,
                        block_type="table" if element.name == "table" else "paragraph",
                    )
                )
        title = soup.title.string if soup.title else None
        return ParsedDocument(blocks=blocks, page_count=1, backend=self.name, title=title)


class DoclingParser(BaseParser):
    """Layout-aware parser. Preserves reading order, tables and section
    hierarchy -- worth the extra CPU on any document with real structure."""

    name = "docling"
    supported = (
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/vnd.openxmlformats-officedocument.presentationml.presentation",
        "text/html",
        "image/png",
        "image/jpeg",
    )

    def __init__(self) -> None:
        self._converter = None

    def _get_converter(self):
        if self._converter is None:
            from docling.document_converter import DocumentConverter  # type: ignore

            self._converter = DocumentConverter()
        return self._converter

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        import tempfile
        from pathlib import Path

        suffix = Path(filename).suffix or ".pdf"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as tmp:
            tmp.write(data)
            tmp.flush()
            result = self._get_converter().convert(tmp.name)

        doc = result.document
        blocks: list[ParsedBlock] = []
        heading_stack: list[str] = []

        for item, level in doc.iterate_items():
            text = getattr(item, "text", None)
            label = str(getattr(item, "label", "paragraph"))
            if not text or not text.strip():
                continue
            page = None
            provenance = getattr(item, "prov", None)
            if provenance:
                page = getattr(provenance[0], "page_no", None)

            if "title" in label or "header" in label or "section" in label:
                heading_stack = heading_stack[:level] + [text.strip()]
                blocks.append(
                    ParsedBlock(
                        text=text.strip(), heading=text.strip(), block_type="heading",
                        page_number=page, section_path=" > ".join(heading_stack),
                    )
                )
            else:
                blocks.append(
                    ParsedBlock(
                        text=text.strip(),
                        page_number=page,
                        heading=heading_stack[-1] if heading_stack else None,
                        section_path=" > ".join(heading_stack) or None,
                        block_type="table" if "table" in label else "paragraph",
                    )
                )

        page_count = len(getattr(doc, "pages", []) or []) or 1
        return ParsedDocument(
            blocks=blocks, page_count=page_count, backend=self.name,
            title=getattr(doc, "name", None),
        )


class ParserRegistry:
    def __init__(self) -> None:
        self._native: list[BaseParser] = [PdfParser(), DocxParser(), HtmlParser(), TextParser()]
        self._docling: DoclingParser | None = None

    def _docling_available(self) -> bool:
        try:
            import docling  # type: ignore  # noqa: F401

            return True
        except ImportError:
            return False

    def select(self, content_type: str, filename: str) -> BaseParser:
        backend = settings.parser_backend
        if backend in {"auto", "docling"} and self._docling_available():
            if self._docling is None:
                self._docling = DoclingParser()
            if self._docling.handles(content_type, filename):
                return self._docling
        for parser in self._native:
            if parser.handles(content_type, filename):
                return parser
        return self._native[-1]

    def parse(self, data: bytes, filename: str, content_type: str) -> ParsedDocument:
        parser = self.select(content_type, filename)
        try:
            return parser.parse(data, filename, content_type)
        except Exception as exc:  # noqa: BLE001
            log.warning(
                "parser.failed_falling_back", backend=parser.name, error=str(exc),
                filename=filename,
            )
            return TextParser().parse(data, filename, "text/plain")


_registry: ParserRegistry | None = None


def get_parser_registry() -> ParserRegistry:
    global _registry
    if _registry is None:
        _registry = ParserRegistry()
    return _registry
